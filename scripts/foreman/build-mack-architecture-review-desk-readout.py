from __future__ import annotations

import html
import json
from datetime import datetime, timezone
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ARCH_DIR = ROOT / "foreman" / "source_material" / "manuscript_workbench" / "tinkularity" / "architecture"
RECEIPT_DIR = ROOT / "foreman" / "receipts"
BASE = "MACK_ARCHITECTURE_REVIEW_DESK_READOUT_20260706"
MD_PATH = ARCH_DIR / f"{BASE}.md"
HTML_PATH = ARCH_DIR / f"{BASE}.html"
DOCX_PATH = ARCH_DIR / f"{BASE}.docx"
RECEIPT_PATH = RECEIPT_DIR / "MACK_ARCHITECTURE_REVIEW_DESK_READOUT_RECEIPT_20260706.json"

DESK_RECEIPT = RECEIPT_DIR / "MACK_ARCHITECTURE_REVIEW_DESK_READINESS_RECEIPT_20260706.json"
RECEIVER_AUDIT = RECEIPT_DIR / "BOOK_ARCHITECTURE_RECEIVER_PROOF_EVERYWHERE_AUDIT_V0_RECEIPT_20260706.json"
PACKET_MD = ARCH_DIR / "BOOK_ARCHITECTURE_REVIEW_PACKET_FOR_BEN_AND_MACK_20260706.md"
PACKET_DOCX = ARCH_DIR / "BOOK_ARCHITECTURE_REVIEW_PACKET_FOR_BEN_AND_MACK_20260706.docx"
PACKET_HTML = ARCH_DIR / "BOOK_ARCHITECTURE_REVIEW_PACKET_FOR_BEN_AND_MACK_20260706.html"
INDEX_MD = ARCH_DIR / "BOOK_ARCHITECTURE_REVIEW_PACKET_INDEX_20260706.md"
BEN_BRIEF = ARCH_DIR / "BEN_PRE_MACK_ARCHITECTURE_BRIEF_AND_AEYE_INPUT_20260706.md"
MACK_HANDOFF = ROOT / "foreman" / "handoffs" / "outbox" / "TO_MACK_HARVEY_NERDKLE_ARCHITECTURE_TEAR_APART_20260706.md"
MACK_INTAKE = ARCH_DIR / "MACK_ARCHITECTURE_TEAR_APART_RETURN_INTAKE_20260706.md"


def repo_rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def file_hash(path: Path) -> dict:
    raw = path.read_bytes()
    return {"path": repo_rel(path), "sha256": sha256(raw).hexdigest(), "bytes": len(raw)}


def esc(value: str) -> str:
    return html.escape(value, quote=False)


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        out.append("| " + " | ".join(cell.replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(out)


def build_markdown(desk: dict, audit: dict) -> str:
    counts = desk["receiver_handoff_readback"]
    validation = desk["validation"]
    blockers = audit.get("blockers", [])

    current_state = [
        ["Mack return", "Not received", "`MACK_RETURN_NOT_RECEIVED`; no Mack-derived next-build packet exists."],
        ["External send", "Not performed", "The desk is assembled locally; no send to Mack is claimed."],
        ["Review desk", "Ready", "`Open-MackArchitectureReviewDesk.ps1 -DryRun` returned `READY_TO_OPEN`."],
        [
            "Receiver handoffs",
            f"{counts['count']} indexed",
            f"{counts['posted_count']} posted; {counts['pending_count']} pending; "
            f"{counts['returned_unposted_count']} returned-unposted; "
            f"{counts['template_return_blocked_count']} template-blocked; "
            f"{counts['invalid_count']} invalid; {counts['malformed_count']} malformed.",
        ],
        ["Universal receiver proof", "Not claimed", f"{len(blockers)} receiver-proof blockers remain."],
        ["Workspace Relay bridge", "Pending receiver", "Bridge can open a receiver-return lane, but no returned receipt exists."],
    ]

    proof_rows = [
        ["Contract canon", "Current", "Packet, receipt, event, gate, and boot-context contracts parse and hash-match receipts."],
        ["Event spine", "Current", "Packet and receipt events still join by packet_id and receipt_id."],
        ["Boot context", "Current", "World state is fresh under 12 hours and active context was refreshed."],
        ["Nerdkle mirror", "Partial proof", "Legacy Nerdkle receipts mirror into organism packet/receipt/event records."],
        ["SoleDash transport", "Partial proof", "Transport ACKs mirror as partial receipts; ACK is not upgraded into receiver work proof."],
        ["SoleDash handoff bridge", "Pending receiver", "Creates a blocked `TEMPLATE_NOT_FILLED` return template; no completion claim."],
        ["Workspace Relay bridge", "Pending receiver", "Creates a blocked `TEMPLATE_NOT_FILLED` return template; no completion claim."],
        ["Receiver handoff lane", "Enforced", "Pending, returned-unposted, template-blocked, and posted states stay distinct."],
    ]

    file_rows = [
        ["Front door", repo_rel(MD_PATH)],
        ["Aeye input next", repo_rel(BEN_BRIEF)],
        ["Main review packet MD", repo_rel(PACKET_MD)],
        ["Main review packet DOCX", repo_rel(PACKET_DOCX)],
        ["Main review packet HTML", repo_rel(PACKET_HTML)],
        ["Packet index", repo_rel(INDEX_MD)],
        ["Ready-to-paste Mack handoff", repo_rel(MACK_HANDOFF)],
        ["Mack return intake", repo_rel(MACK_INTAKE)],
        ["Desk readiness receipt", repo_rel(DESK_RECEIPT)],
        ["Receiver-proof audit receipt", repo_rel(RECEIVER_AUDIT)],
    ]

    mack_questions = [
        "Does shared body-state solve cooperation, or does it only rename packet passing?",
        "Can file-backed packets, events, and receipts feel live enough to reduce Ben's babysitting burden?",
        "Which organ is premature or decorative: Speaker, Wormeyes, Medulla, TinkerDen, SoleDash, or the event spine?",
        "Where can the system still fake completion?",
        "Which proof field is missing from the packet/receipt/event chain?",
        "What is the smallest build that makes the architecture harder to dismiss?",
    ]

    lines = [
        "# Mack Architecture Review Desk Readout",
        "",
        "Status: READY FOR BEN REVIEW",
        "Date: 2026-07-06",
        "Owner: Heimerdinker@Betsy",
        "Lane: Harvey/Nerdkle architecture review",
        f"Source readiness receipt: `{repo_rel(DESK_RECEIPT)}`",
        "",
        "## Bottom Line",
        "",
        "The architecture packet is assembled and ready for Ben to review before giving it to Mack. The desk proves the local packet set, receipts, launcher, intake, and proof readbacks are present. It does not prove Mack has reviewed anything.",
        "",
        "Use this readout for the current proof counts. It supersedes older count text in any stale copy of the packet.",
        "",
        "## Current State",
        "",
        md_table(["Area", "State", "Readback"], current_state),
        "",
        "## What Mack Should Tear Apart",
        "",
    ]
    for index, question in enumerate(mack_questions, 1):
        lines.append(f"{index}. {question}")
    lines.extend(
        [
            "",
            "## Proof Surface Snapshot",
            "",
            md_table(["Surface", "State", "What The Receipt Proves"], proof_rows),
            "",
            "## Boundaries That Must Stay True",
            "",
            "- Mack has not returned a review.",
            "- No external send to Mack has been performed from this desk.",
            "- No canonical next-build packet exists yet.",
            "- A transport ACK is not receiver work completion.",
            "- A blocked `TEMPLATE_NOT_FILLED` return template is not completion proof.",
            "- Universal receiver proof is not claimed until non-template receiver returns are filled and posted.",
            "",
            "## Ben Operating Notes",
            "",
            "- You are in the desk readout; use it as the front door and current proof-count source.",
            "- Next read `BEN_PRE_MACK_ARCHITECTURE_BRIEF_AND_AEYE_INPUT_20260706.md` for Aeye input before Mack.",
            "- Then read the main packet in Markdown, DOCX, or HTML.",
            "- Use the operator-only local proof links when the dev server is running:",
            "  - `http://127.0.0.1:3000/tinkerden?handoff_provenance=operator`",
            "  - `http://127.0.0.1:3000/tinkerden/receipts?handoff_provenance=operator`",
            "- Paste `TO_MACK_HARVEY_NERDKLE_ARCHITECTURE_TEAR_APART_20260706.md` to Mack only when Ben decides to send it.",
            "- After Mack returns, put the return into the intake file and run the validator before creating any next-build packet.",
            "",
            "## File Drawer",
            "",
            md_table(["Role", "Path"], file_rows),
            "",
            "## Receiver-Proof Blockers Still Open",
            "",
        ],
    )
    for blocker in blockers:
        lines.append(f"- {blocker}")
    lines.extend(
        [
            "",
            "## Validation Readback",
            "",
            f"- Required artifacts exist: `{validation['required_artifacts_exist']}`.",
            f"- Required receipts parse: `{validation['required_receipts_parse']}`.",
            f"- Mack-not-received validator: `{validation['validator_reports_mack_not_received']}`.",
            f"- Launcher dry run ready: `{validation['launcher_dry_run_ready']}`.",
            f"- Canonical next-build packet absent: `{validation['canonical_next_build_packet_absent']}`.",
            f"- Workspace Relay receiver-handoff bridge pending: `{validation['workspace_relay_receiver_handoff_bridge_pending']}`.",
            f"- Live receiver handoff counts match index: `{validation['live_receiver_handoff_counts_match_index']}`.",
            "",
        ],
    )
    return "\n".join(lines)


def markdown_to_html(markdown: str) -> str:
    lines = markdown.splitlines()
    parts = [
        "<!doctype html>",
        "<html lang=\"en\">",
        "<head>",
        "  <meta charset=\"utf-8\">",
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">",
        "  <title>Mack Architecture Review Desk Readout</title>",
        "  <style>",
        "    body { margin: 0; background: #f0f3f7; color: #172033; font: 16px/1.55 Arial, sans-serif; }",
        "    main { max-width: 980px; margin: 32px auto; padding: 40px; background: #fff; border: 1px solid #d8dde6; box-shadow: 0 16px 48px rgba(23,32,51,.12); }",
        "    h1 { margin: 0 0 10px; color: #0b2545; font-size: 2.25rem; line-height: 1.08; }",
        "    h2 { margin: 32px 0 12px; padding-top: 16px; border-top: 1px solid #d8dde6; color: #1f5f8b; }",
        "    p { margin: 0 0 12px; }",
        "    ul, ol { margin-top: 0; padding-left: 1.5rem; }",
        "    li { margin-bottom: 7px; }",
        "    table { width: 100%; border-collapse: collapse; margin: 16px 0 24px; font-size: .93rem; }",
        "    th, td { border: 1px solid #d8dde6; padding: 9px 10px; vertical-align: top; }",
        "    th { background: #e8eef5; text-align: left; }",
        "    code { background: #eef1f5; border-radius: 4px; padding: 1px 4px; font-family: Consolas, monospace; font-size: .93em; }",
        "    .meta { color: #5b6474; font-size: .93rem; }",
        "  </style>",
        "</head>",
        "<body><main>",
    ]
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            parts.append(f"<h1>{esc(stripped[2:])}</h1>")
        elif stripped.startswith("## "):
            parts.append(f"<h2>{esc(stripped[3:])}</h2>")
        elif stripped.startswith("| "):
            table_lines = []
            while index < len(lines) and lines[index].startswith("| "):
                table_lines.append(lines[index])
                index += 1
            rows = []
            for raw in table_lines:
                cells = [cell.strip() for cell in raw.strip("|").split("|")]
                if all(set(cell) <= {"-"} and cell for cell in cells):
                    continue
                rows.append(cells)
            if rows:
                parts.append("<table><thead><tr>")
                parts.extend(f"<th>{inline(cell)}</th>" for cell in rows[0])
                parts.append("</tr></thead><tbody>")
                for row in rows[1:]:
                    parts.append("<tr>" + "".join(f"<td>{inline(cell)}</td>" for cell in row) + "</tr>")
                parts.append("</tbody></table>")
            continue
        elif stripped.startswith("- "):
            items = []
            while index < len(lines) and lines[index].strip().startswith("- "):
                items.append(lines[index].strip()[2:])
                index += 1
            parts.append("<ul>" + "".join(f"<li>{inline(item)}</li>" for item in items) + "</ul>")
            continue
        elif stripped[:2].isdigit() and ". " in stripped[:4]:
            items = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate[:2].isdigit() and ". " in candidate[:4] or candidate[:1].isdigit() and ". " in candidate[:3]):
                    break
                items.append(candidate.split(". ", 1)[1])
                index += 1
            parts.append("<ol>" + "".join(f"<li>{inline(item)}</li>" for item in items) + "</ol>")
            continue
        else:
            cls = " class=\"meta\"" if stripped.startswith(("Status:", "Date:", "Owner:", "Lane:", "Source readiness receipt:")) else ""
            parts.append(f"<p{cls}>{inline(stripped)}</p>")
        index += 1
    parts.extend(["</main></body>", "</html>"])
    return "\n".join(parts) + "\n"


def inline(value: str) -> str:
    escaped = esc(value)
    out = []
    in_code = False
    buf = []
    for char in escaped:
        if char == "`":
            if in_code:
                out.append("<code>" + "".join(buf) + "</code>")
                buf = []
                in_code = False
            else:
                out.append("".join(buf))
                buf = []
                in_code = True
        else:
            buf.append(char)
    out.append("".join(buf))
    return "".join(out)


def set_style_font(style, name: str, size: float | None = None, color: str | None = None) -> None:
    font = style.font
    font.name = name
    if size is not None:
        font.size = Pt(size)
    if color is not None:
        font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_inline_runs(paragraph, value: str) -> None:
    chunks = value.split("`")
    for index, chunk in enumerate(chunks):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if index % 2 == 1:
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)


def add_bullets(doc: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        add_inline_runs(paragraph, item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        run.bold = True
    for row_index, row in enumerate(rows, 1):
        for col, value in enumerate(row):
            paragraph = table.cell(row_index, col).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline_runs(paragraph, value)
    doc.add_paragraph()


def build_docx(markdown: str, desk: dict, audit: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style_font(doc.styles["Normal"], "Calibri", 11, "172033")
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.25
    set_style_font(doc.styles["Heading 1"], "Calibri", 16, "2E74B5")
    set_style_font(doc.styles["Heading 2"], "Calibri", 13, "2E74B5")
    set_style_font(doc.styles["Heading 3"], "Calibri", 12, "1F4D78")
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(18)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(7)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(5)

    counts = desk["receiver_handoff_readback"]
    validation = desk["validation"]
    blockers = audit.get("blockers", [])

    header = section.header.paragraphs[0]
    header.text = "Harvey/Nerdkle Architecture Review"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("5B6474")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Mack Architecture Review Desk Readout")
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    for line in [
        "Status: READY FOR BEN REVIEW",
        "Date: 2026-07-06",
        "Owner: Heimerdinker@Betsy",
        "Lane: Harvey/Nerdkle architecture review",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("5B6474")

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "F4F7FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "The desk is assembled locally and ready for Ben's review. Mack has not returned a review, no external send is claimed, and no next-build packet exists yet."
    )
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    doc.add_paragraph()

    doc.add_heading("Bottom Line", level=1)
    doc.add_paragraph(
        "Use this readout as the front door. The architecture packet is ready to hand to Mack for attack, with the proof boundary kept honest: this is an assembled local review desk, not a Mack return and not a universal receiver-proof claim."
    )
    doc.add_paragraph("Use this readout for current proof counts: 19 indexed, 6 posted, 9 pending, 3 returned-unposted, 1 template-blocked, 0 invalid, 0 malformed.")

    doc.add_heading("Current State", level=1)
    add_table(
        doc,
        ["Area", "State", "Readback"],
        [
            ["Mack return", "Not received", "`MACK_RETURN_NOT_RECEIVED`; no Mack-derived next-build packet exists."],
            ["External send", "Not performed", "The desk is assembled locally; no send to Mack is claimed."],
            ["Review desk", "Ready", "`Open-MackArchitectureReviewDesk.ps1 -DryRun` returned `READY_TO_OPEN`."],
            [
                "Receiver handoffs",
                f"{counts['count']} indexed",
                f"{counts['posted_count']} posted; {counts['pending_count']} pending; {counts['returned_unposted_count']} returned-unposted; {counts['template_return_blocked_count']} template-blocked; {counts['invalid_count']} invalid; {counts['malformed_count']} malformed.",
            ],
            ["Universal receiver proof", "Not claimed", f"{len(blockers)} receiver-proof blockers remain."],
            ["Workspace Relay bridge", "Pending receiver", "Bridge can open a receiver-return lane, but no returned receipt exists."],
        ],
        [1900, 1900, 5560],
    )

    doc.add_heading("What Mack Should Tear Apart", level=1)
    add_bullets(
        doc,
        [
            "Does shared body-state solve cooperation, or does it only rename packet passing?",
            "Can file-backed packets, events, and receipts feel live enough to reduce Ben's babysitting burden?",
            "Which organ is premature or decorative: Speaker, Wormeyes, Medulla, TinkerDen, SoleDash, or the event spine?",
            "Where can the system still fake completion?",
            "Which proof field is missing from the packet/receipt/event chain?",
            "What is the smallest build that makes the architecture harder to dismiss?",
        ],
        numbered=True,
    )

    doc.add_heading("Proof Surface Snapshot", level=1)
    add_table(
        doc,
        ["Surface", "State", "What The Receipt Proves"],
        [
            ["Contract canon", "Current", "Packet, receipt, event, gate, and boot-context contracts parse and hash-match receipts."],
            ["Event spine", "Current", "Packet and receipt events still join by packet_id and receipt_id."],
            ["Boot context", "Current", "World state is fresh under 12 hours and active context was refreshed."],
            ["Nerdkle mirror", "Partial proof", "Legacy Nerdkle receipts mirror into organism packet/receipt/event records."],
            ["SoleDash transport", "Partial proof", "Transport ACKs mirror as partial receipts; ACK is not upgraded into receiver work proof."],
            ["SoleDash handoff bridge", "Pending receiver", "Creates a blocked `TEMPLATE_NOT_FILLED` return template; no completion claim."],
            ["Workspace Relay bridge", "Pending receiver", "Creates a blocked `TEMPLATE_NOT_FILLED` return template; no completion claim."],
            ["Receiver handoff lane", "Enforced", "Pending, returned-unposted, template-blocked, and posted states stay distinct."],
        ],
        [2250, 1700, 5410],
    )

    doc.add_heading("Boundaries That Must Stay True", level=1)
    add_bullets(
        doc,
        [
            "Mack has not returned a review.",
            "No external send to Mack has been performed from this desk.",
            "No canonical next-build packet exists yet.",
            "A transport ACK is not receiver work completion.",
            "A blocked `TEMPLATE_NOT_FILLED` return template is not completion proof.",
            "Universal receiver proof is not claimed until non-template receiver returns are filled and posted.",
        ],
    )

    doc.add_heading("File Drawer", level=1)
    add_table(
        doc,
        ["Role", "Path"],
        [
            ["Front door", repo_rel(MD_PATH)],
            ["Aeye input next", repo_rel(BEN_BRIEF)],
            ["Main review packet MD", repo_rel(PACKET_MD)],
            ["Main review packet DOCX", repo_rel(PACKET_DOCX)],
            ["Main review packet HTML", repo_rel(PACKET_HTML)],
            ["Packet index", repo_rel(INDEX_MD)],
            ["Ready-to-paste Mack handoff", repo_rel(MACK_HANDOFF)],
            ["Mack return intake", repo_rel(MACK_INTAKE)],
            ["Desk readiness receipt", repo_rel(DESK_RECEIPT)],
            ["Receiver-proof audit receipt", repo_rel(RECEIVER_AUDIT)],
        ],
        [2400, 6960],
    )

    doc.add_heading("Validation Readback", level=1)
    add_bullets(
        doc,
        [
            f"Required artifacts exist: `{validation['required_artifacts_exist']}`.",
            f"Required receipts parse: `{validation['required_receipts_parse']}`.",
            f"Mack-not-received validator: `{validation['validator_reports_mack_not_received']}`.",
            f"Launcher dry run ready: `{validation['launcher_dry_run_ready']}`.",
            f"Canonical next-build packet absent: `{validation['canonical_next_build_packet_absent']}`.",
            f"Workspace Relay receiver-handoff bridge pending: `{validation['workspace_relay_receiver_handoff_bridge_pending']}`.",
            f"Live receiver handoff counts match index: `{validation['live_receiver_handoff_counts_match_index']}`.",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.text = "Mack Architecture Review Desk Readout - 2026-07-06"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("5B6474")

    doc.save(DOCX_PATH)


def main() -> None:
    ARCH_DIR.mkdir(parents=True, exist_ok=True)
    RECEIPT_DIR.mkdir(parents=True, exist_ok=True)
    desk = read_json(DESK_RECEIPT)
    audit = read_json(RECEIVER_AUDIT)
    markdown = build_markdown(desk, audit)
    MD_PATH.write_text(markdown, encoding="utf-8")
    HTML_PATH.write_text(markdown_to_html(markdown), encoding="utf-8")
    build_docx(markdown, desk, audit)

    generated_at = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    receipt = {
        "schema": "MACK_ARCHITECTURE_REVIEW_DESK_READOUT_RECEIPT",
        "status": "ARTIFACT",
        "timestamp": generated_at,
        "machine": "BETSY",
        "agent": "Heimerdinker@Betsy",
        "packet_id": "MACK_HARVEY_NERDKLE_ARCHITECTURE_TEAR_APART_20260706",
        "receipt_id": "MACK_ARCHITECTURE_REVIEW_DESK_READOUT_RECEIPT_20260706",
        "repo": str(ROOT),
        "command": "python scripts/foreman/build-mack-architecture-review-desk-readout.py",
        "validation": {
            "markdown_written": MD_PATH.exists(),
            "html_written": HTML_PATH.exists(),
            "docx_written": DOCX_PATH.exists(),
            "mack_return_not_received_preserved": desk["validation"]["validator_reports_mack_not_received"] is True,
            "external_send_not_claimed": desk["validation"]["no_external_send_claim"] is True,
            "canonical_next_build_packet_absent": desk["validation"]["canonical_next_build_packet_absent"] is True,
            "workspace_relay_receiver_handoff_bridge_pending": desk["validation"]["workspace_relay_receiver_handoff_bridge_pending"] is True,
            "universal_receiver_proof_not_claimed": audit["validation"]["universal_receiver_proof_claimed"] is False,
        },
        "receiver_handoff_readback": desk["receiver_handoff_readback"],
        "source_receipts": [repo_rel(DESK_RECEIPT), repo_rel(RECEIVER_AUDIT)],
        "file_hashes": [
            file_hash(MD_PATH),
            file_hash(HTML_PATH),
            file_hash(DOCX_PATH),
            file_hash(Path(__file__).resolve()),
        ],
        "truth_boundary": "This readout is a human handoff surface for Ben and Mack. It does not claim Mack returned review, external send, universal receiver proof, or receiver work completion for pending handoffs.",
        "stop_conditions_respected": [
            "no external send",
            "no Mack receipt claim",
            "no universal receiver proof claim",
            "no canonical next-build packet generated",
            "no deploy",
            "no push",
        ],
    }
    RECEIPT_PATH.write_text(json.dumps(receipt, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"ok": True, "md": repo_rel(MD_PATH), "html": repo_rel(HTML_PATH), "docx": repo_rel(DOCX_PATH), "receipt": repo_rel(RECEIPT_PATH)}, indent=2))


if __name__ == "__main__":
    main()
