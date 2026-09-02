from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ARCH_DIR = ROOT / "foreman" / "source_material" / "manuscript_workbench" / "tinkularity" / "architecture"
BASE = "BOOK_ARCHITECTURE_REVIEW_PACKET_FOR_BEN_AND_MACK_20260706"
MD_PATH = ARCH_DIR / f"{BASE}.md"
HTML_PATH = ARCH_DIR / f"{BASE}.html"
DOCX_PATH = ARCH_DIR / f"{BASE}.docx"
FRONT_MATTER_PREFIXES = (
    "Status:",
    "Date:",
    "Prepared by:",
    "Prepared for:",
    "Source chapter:",
    "Primary stack lock:",
)


def parse_markdown(text: str) -> list[dict]:
    blocks: list[dict] = []
    lines = text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            lang = stripped.strip("`").strip()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            blocks.append({"type": "heading", "level": len(heading.group(1)), "text": heading.group(2)})
            index += 1
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith("> "):
                quote_lines.append(lines[index].strip()[2:])
                index += 1
            blocks.append({"type": "quote", "text": " ".join(quote_lines)})
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = []
            for raw in table_lines:
                cells = [cell.strip() for cell in raw.strip("|").split("|")]
                if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            blocks.append({"type": "table", "rows": rows})
            continue

        if re.match(r"^[-*]\s+", stripped):
            items = []
            while index < len(lines) and re.match(r"^[-*]\s+", lines[index].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[index].strip()))
                index += 1
            blocks.append({"type": "ul", "items": items})
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while index < len(lines) and re.match(r"^\d+\.\s+", lines[index].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[index].strip()))
                index += 1
            blocks.append({"type": "ol", "items": items})
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line:
                break
            if paragraph_lines[0].startswith(FRONT_MATTER_PREFIXES) and next_line.startswith(FRONT_MATTER_PREFIXES):
                break
            if next_line.startswith("```") or next_line.startswith("|") or next_line.startswith("> "):
                break
            if re.match(r"^(#{1,6})\s+", next_line) or re.match(r"^[-*]\s+", next_line) or re.match(r"^\d+\.\s+", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        blocks.append({"type": "p", "text": " ".join(paragraph_lines)})

    return blocks


def inline_html(value: str) -> str:
    escaped = html.escape(value)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    return escaped


def build_html(blocks: list[dict]) -> None:
    parts = [
        "<!doctype html>",
        "<html lang=\"en\">",
        "<head>",
        "  <meta charset=\"utf-8\">",
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">",
        "  <title>Harvey/Nerdkle Architecture Review Packet</title>",
        "  <style>",
        "    :root { color-scheme: light; --ink: #172033; --muted: #5b6474; --line: #d8dde6; --accent: #1f5f8b; --soft: #f4f7fb; }",
        "    body { margin: 0; background: #eef2f7; color: var(--ink); font: 16px/1.58 Arial, sans-serif; }",
        "    main { max-width: 980px; margin: 32px auto; padding: 42px; background: #fff; border: 1px solid var(--line); box-shadow: 0 18px 55px rgba(23, 32, 51, .12); }",
        "    h1 { margin: 0 0 8px; font-size: 2.3rem; line-height: 1.08; color: #0b2545; }",
        "    h2 { margin-top: 34px; padding-top: 18px; border-top: 1px solid var(--line); color: var(--accent); }",
        "    h3 { margin-top: 24px; color: #1f4d78; }",
        "    p { margin: 0 0 13px; }",
        "    ul, ol { margin-top: 0; padding-left: 1.55rem; }",
        "    li { margin-bottom: 6px; }",
        "    blockquote { margin: 18px 0; padding: 16px 20px; border-left: 5px solid var(--accent); background: var(--soft); font-weight: 700; }",
        "    table { width: 100%; border-collapse: collapse; margin: 18px 0 24px; font-size: .92rem; }",
        "    th, td { border: 1px solid var(--line); padding: 9px 10px; vertical-align: top; }",
        "    th { background: #f2f4f7; text-align: left; }",
        "    code { padding: 1px 4px; border-radius: 4px; background: #eef1f5; font-family: Consolas, monospace; font-size: .93em; }",
        "    pre { overflow-x: auto; padding: 16px; border: 1px solid var(--line); border-radius: 8px; background: #111827; color: #eef2ff; white-space: pre-wrap; }",
        "    pre code { padding: 0; background: transparent; color: inherit; }",
        "    .small { color: var(--muted); font-size: .92rem; }",
        "  </style>",
        "</head>",
        "<body>",
        "<main>",
    ]
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = block["level"]
            text = inline_html(block["text"])
            parts.append(f"<h{level}>{text}</h{level}>")
        elif kind == "p":
            cls = " class=\"small\"" if block["text"].startswith(("Status:", "Date:", "Prepared")) else ""
            parts.append(f"<p{cls}>{inline_html(block['text'])}</p>")
        elif kind == "quote":
            parts.append(f"<blockquote>{inline_html(block['text'])}</blockquote>")
        elif kind in {"ul", "ol"}:
            tag = kind
            parts.append(f"<{tag}>")
            for item in block["items"]:
                parts.append(f"  <li>{inline_html(item)}</li>")
            parts.append(f"</{tag}>")
        elif kind == "code":
            lang = html.escape(block.get("lang", ""))
            parts.append(f"<pre><code class=\"language-{lang}\">{html.escape(block['text'])}</code></pre>")
        elif kind == "table":
            rows = block["rows"]
            if not rows:
                continue
            parts.append("<table>")
            parts.append("<thead><tr>" + "".join(f"<th>{inline_html(cell)}</th>" for cell in rows[0]) + "</tr></thead>")
            parts.append("<tbody>")
            for row in rows[1:]:
                parts.append("<tr>" + "".join(f"<td>{inline_html(cell)}</td>" for cell in row) + "</tr>")
            parts.append("</tbody></table>")
    parts.extend(["</main>", "</body>", "</html>"])
    HTML_PATH.write_text("\n".join(parts) + "\n", encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def add_inline_runs(paragraph, value: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(value):
        if match.start() > pos:
            paragraph.add_run(value[pos:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        pos = match.end()
    if pos < len(value):
        paragraph.add_run(value[pos:])


def set_style_font(style, name: str, size: float | None = None, color: str | None = None) -> None:
    font = style.font
    font.name = name
    if size is not None:
        font.size = Pt(size)
    if color is not None:
        font.color.rgb = RGBColor.from_string(color)


def build_docx(blocks: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    set_style_font(doc.styles["Normal"], "Calibri", 11, "172033")
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.1
    set_style_font(doc.styles["Heading 1"], "Calibri", 16, "2E74B5")
    set_style_font(doc.styles["Heading 2"], "Calibri", 13, "2E74B5")
    set_style_font(doc.styles["Heading 3"], "Calibri", 12, "1F4D78")
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(4)

    first = True
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = block["level"]
            if first and level == 1:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(4)
                run = paragraph.add_run(block["text"])
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor.from_string("0B2545")
                first = False
            else:
                doc.add_heading(block["text"], level=min(level, 3))
        elif kind == "p":
            paragraph = doc.add_paragraph()
            if block["text"].startswith(("Status:", "Date:", "Prepared", "Source chapter:", "Primary stack lock:")):
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(block["text"])
                run.font.color.rgb = RGBColor.from_string("5B6474")
                run.font.size = Pt(9.5)
            else:
                add_inline_runs(paragraph, block["text"])
        elif kind == "quote":
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            set_cell_margins(table)
            cell = table.cell(0, 0)
            set_cell_width(cell, 6.3)
            set_cell_shading(cell, "F4F7FB")
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(block["text"])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("0B2545")
        elif kind in {"ul", "ol"}:
            style = "List Bullet" if kind == "ul" else "List Number"
            for item in block["items"]:
                paragraph = doc.add_paragraph(style=style)
                add_inline_runs(paragraph, item)
        elif kind == "code":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("172033")
        elif kind == "table":
            rows = block["rows"]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = "Table Grid"
            table.autofit = False
            set_cell_margins(table)
            col_count = max(len(row) for row in rows)
            widths = [6.5 / col_count] * col_count
            if col_count == 4:
                widths = [1.2, 1.55, 1.65, 2.1]
            for row_index, row in enumerate(rows):
                cells = table.rows[row_index].cells
                for col_index, cell in enumerate(cells):
                    set_cell_width(cell, widths[col_index])
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(0)
                    text = row[col_index] if col_index < len(row) else ""
                    add_inline_runs(paragraph, text)
                    if row_index == 0:
                        set_cell_shading(cell, "F2F4F7")
                        for run in paragraph.runs:
                            run.bold = True

    footer = section.footer.paragraphs[0]
    footer.text = "Harvey/Nerdkle Architecture Review Packet - V0.2 - 2026-07-06"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("5B6474")

    doc.save(DOCX_PATH)


def main() -> None:
    markdown = MD_PATH.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown)
    build_html(blocks)
    build_docx(blocks)
    print(f"wrote {HTML_PATH.relative_to(ROOT)}")
    print(f"wrote {DOCX_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
