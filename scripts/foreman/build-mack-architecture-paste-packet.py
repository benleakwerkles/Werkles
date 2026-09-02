from __future__ import annotations

import html
import json
import re
from datetime import datetime, timezone
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ARCH_DIR = ROOT / "foreman" / "source_material" / "manuscript_workbench" / "tinkularity" / "architecture"
RECEIPT_DIR = ROOT / "foreman" / "receipts"
HANDOFF = ROOT / "foreman" / "handoffs" / "outbox" / "TO_MACK_HARVEY_NERDKLE_ARCHITECTURE_TEAR_APART_20260706.md"
DESK_RECEIPT = RECEIPT_DIR / "MACK_ARCHITECTURE_REVIEW_DESK_READINESS_RECEIPT_20260706.json"

BASE = "MACK_ARCHITECTURE_TEAR_APART_PASTE_PACKET_20260706"
MD_PATH = ARCH_DIR / f"{BASE}.md"
HTML_PATH = ARCH_DIR / f"{BASE}.html"
DOCX_PATH = ARCH_DIR / f"{BASE}.docx"
RECEIPT_PATH = RECEIPT_DIR / f"{BASE}_RECEIPT.json"

READOUT_MD = ARCH_DIR / "MACK_ARCHITECTURE_REVIEW_DESK_READOUT_20260706.md"
READOUT_DOCX = ARCH_DIR / "MACK_ARCHITECTURE_REVIEW_DESK_READOUT_20260706.docx"
READOUT_HTML = ARCH_DIR / "MACK_ARCHITECTURE_REVIEW_DESK_READOUT_20260706.html"
BEN_BRIEF = ARCH_DIR / "BEN_PRE_MACK_ARCHITECTURE_BRIEF_AND_AEYE_INPUT_20260706.md"
CONNECTION_MD = ARCH_DIR / "BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.md"
CONNECTION_HTML = ARCH_DIR / "BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.html"
CONNECTION_JSON = ARCH_DIR / "BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.json"
CONNECTION_MERMAID = ARCH_DIR / "BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.mmd"
SCORECARD_MD = ARCH_DIR / "MACK_ARCHITECTURE_ATTACK_SCORECARD_20260706.md"
SCORECARD_HTML = ARCH_DIR / "MACK_ARCHITECTURE_ATTACK_SCORECARD_20260706.html"
SCORECARD_JSON = ARCH_DIR / "MACK_ARCHITECTURE_ATTACK_SCORECARD_20260706.json"
PACKET_MD = ARCH_DIR / "BOOK_ARCHITECTURE_REVIEW_PACKET_FOR_BEN_AND_MACK_20260706.md"
PACKET_DOCX = ARCH_DIR / "BOOK_ARCHITECTURE_REVIEW_PACKET_FOR_BEN_AND_MACK_20260706.docx"
PACKET_HTML = ARCH_DIR / "BOOK_ARCHITECTURE_REVIEW_PACKET_FOR_BEN_AND_MACK_20260706.html"
RETURN_INTAKE = ARCH_DIR / "MACK_ARCHITECTURE_TEAR_APART_RETURN_INTAKE_20260706.md"
SCORECARD_INTAKE = ARCH_DIR / "MACK_ARCHITECTURE_ATTACK_SCORECARD_RETURN_INTAKE_20260706.md"
RETURN_TEMPLATE = ARCH_DIR / "MACK_ARCHITECTURE_TEAR_APART_RETURN_TEMPLATE_20260706.md"


def repo_rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def file_hash(path: Path) -> dict:
    raw = path.read_bytes()
    return {"path": repo_rel(path), "sha256": sha256(raw).hexdigest(), "bytes": len(raw)}


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def extract_paste_block(handoff: str) -> str:
    match = re.search(r"```text\s*(.*?)\s*```", handoff, re.DOTALL)
    if not match:
        raise ValueError("Mack handoff is missing the fenced text paste block.")
    return match.group(1).strip()


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    output = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        output.append("| " + " | ".join(cell.replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(output)


def build_markdown(handoff: str, desk: dict) -> str:
    paste_block = extract_paste_block(handoff)
    counts = desk["receiver_handoff_readback"]
    artifact_rows = [
        ["Desk readout MD", repo_rel(READOUT_MD)],
        ["Desk readout DOCX", repo_rel(READOUT_DOCX)],
        ["Desk readout HTML", repo_rel(READOUT_HTML)],
        ["Ben pre-Mack Aeye brief", repo_rel(BEN_BRIEF)],
        ["Connection map MD", repo_rel(CONNECTION_MD)],
        ["Connection map HTML", repo_rel(CONNECTION_HTML)],
        ["Connection map JSON", repo_rel(CONNECTION_JSON)],
        ["Connection map Mermaid", repo_rel(CONNECTION_MERMAID)],
        ["Mack attack scorecard MD", repo_rel(SCORECARD_MD)],
        ["Mack attack scorecard HTML", repo_rel(SCORECARD_HTML)],
        ["Mack attack scorecard JSON", repo_rel(SCORECARD_JSON)],
        ["Main packet MD", repo_rel(PACKET_MD)],
        ["Main packet DOCX", repo_rel(PACKET_DOCX)],
        ["Main packet HTML", repo_rel(PACKET_HTML)],
        ["Ready-to-paste handoff source", repo_rel(HANDOFF)],
        ["Return template", repo_rel(RETURN_TEMPLATE)],
        ["Return intake", repo_rel(RETURN_INTAKE)],
        ["Scorecard return intake", repo_rel(SCORECARD_INTAKE)],
    ]

    lines = [
        "# Mack Architecture Tear-Apart Paste Packet",
        "",
        "Status: READY_TO_COPY_NOT_SENT",
        "Date: 2026-07-06",
        "Owner: Heimerdinker@Betsy",
        "Lane: Harvey/Nerdkle architecture review",
        f"Source handoff: `{repo_rel(HANDOFF)}`",
        "",
        "## Purpose",
        "",
        "This packet is the single human-facing copy surface for giving Mack the Harvey/Nerdkle architecture review mission. It is generated from the live ready-to-paste handoff so Ben can review the context, copy the Mack block, and later capture Mack's return without inventing success.",
        "",
        "## Read Order For Ben",
        "",
        "1. Read `MACK_ARCHITECTURE_REVIEW_DESK_READOUT_20260706.md` first.",
        "2. Read `BEN_PRE_MACK_ARCHITECTURE_BRIEF_AND_AEYE_INPUT_20260706.md` next.",
        "3. Read `BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.md` third.",
        "4. Read `MACK_ARCHITECTURE_ATTACK_SCORECARD_20260706.md` fourth.",
        "5. Read the main architecture packet in Markdown, DOCX, or HTML.",
        "6. If Ben decides to send, copy the Mack block below to Mack.",
        "7. Paste Mack's returned `MACK REVIEW RETURN` block into the review intake and `MACK SCORECARD RETURN` into the scorecard intake, then run the validators.",
        "",
        "## Current Proof Counts",
        "",
        f"- Receiver handoffs indexed: {counts['count']}.",
        f"- Posted: {counts['posted_count']}.",
        f"- Pending receiver: {counts['pending_count']}.",
        f"- Returned-unposted: {counts['returned_unposted_count']}.",
        f"- Template-return blocked: {counts['template_return_blocked_count']}.",
        f"- Invalid: {counts['invalid_count']}.",
        f"- Malformed: {counts['malformed_count']}.",
        "",
        "## Truth Boundary",
        "",
        "- This packet does not send anything to Mack.",
        "- Mack has not returned a review.",
        "- No Mack receipt is claimed.",
        "- No canonical next-build packet exists.",
        "- Universal receiver proof is not claimed while pending receiver-return lanes remain.",
        "- A transport ACK or blocked template is not receiver work completion.",
        "",
        "## Copy/Paste Block For Mack",
        "",
        "```text",
        paste_block,
        "```",
        "",
        "## After Mack Returns",
        "",
        "Paste Mack's returned `MACK REVIEW RETURN` block into:",
        "",
        f"`{repo_rel(RETURN_INTAKE)}`",
        "",
        "Then run:",
        "",
        "```powershell",
        "node scripts\\foreman\\mack-architecture-return-intake-validator.mjs",
        "```",
        "",
        "Paste Mack's returned `MACK SCORECARD RETURN` block into:",
        "",
        f"`{repo_rel(SCORECARD_INTAKE)}`",
        "",
        "Then run:",
        "",
        "```powershell",
        "node scripts\\foreman\\mack-architecture-scorecard-return-validator.mjs",
        "```",
        "",
        "Only after Mack returns a complete structured block and Ben explicitly accepts the direction should the guarded conversion path be used:",
        "",
        "```powershell",
        "node scripts\\foreman\\mack-architecture-return-intake-validator.mjs --convert --ben-accepted",
        "```",
        "",
        "## Artifact Drawer",
        "",
        md_table(["Role", "Path"], artifact_rows),
        "",
    ]
    return "\n".join(lines)


def inline_html(value: str) -> str:
    escaped = html.escape(value)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    return escaped


def markdown_to_html(markdown: str) -> str:
    lines = markdown.splitlines()
    parts = [
        "<!doctype html>",
        "<html lang=\"en\">",
        "<head>",
        "  <meta charset=\"utf-8\">",
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">",
        "  <title>Mack Architecture Tear-Apart Paste Packet</title>",
        "  <style>",
        "    body { margin: 0; background: #eef2f7; color: #172033; font: 16px/1.56 Arial, sans-serif; }",
        "    main { max-width: 980px; margin: 32px auto; padding: 42px; background: #fff; border: 1px solid #d8dde6; box-shadow: 0 18px 55px rgba(23,32,51,.12); }",
        "    h1 { margin: 0 0 10px; color: #0b2545; font-size: 2.25rem; line-height: 1.08; }",
        "    h2 { margin: 32px 0 12px; padding-top: 16px; border-top: 1px solid #d8dde6; color: #1f5f8b; }",
        "    p { margin: 0 0 12px; }",
        "    ul, ol { margin-top: 0; padding-left: 1.5rem; }",
        "    li { margin-bottom: 7px; }",
        "    table { width: 100%; border-collapse: collapse; margin: 16px 0 24px; font-size: .93rem; }",
        "    th, td { border: 1px solid #d8dde6; padding: 9px 10px; vertical-align: top; }",
        "    th { background: #e8eef5; text-align: left; }",
        "    code { background: #eef1f5; border-radius: 4px; padding: 1px 4px; font-family: Consolas, monospace; font-size: .93em; }",
        "    pre { overflow-x: auto; padding: 16px; border: 1px solid #d8dde6; border-radius: 6px; background: #111827; color: #eef2ff; white-space: pre-wrap; }",
        "    pre code { background: transparent; color: inherit; padding: 0; }",
        "    .meta { color: #5b6474; font-size: .93rem; }",
        "  </style>",
        "</head>",
        "<body><main>",
    ]
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            lang = stripped.strip("`").strip()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            parts.append(f"<pre><code class=\"language-{html.escape(lang)}\">{html.escape(chr(10).join(code_lines))}</code></pre>")
            continue
        if stripped.startswith("# "):
            parts.append(f"<h1>{inline_html(stripped[2:])}</h1>")
        elif stripped.startswith("## "):
            parts.append(f"<h2>{inline_html(stripped[3:])}</h2>")
        elif stripped.startswith("| "):
            table_lines = []
            while index < len(lines) and lines[index].startswith("| "):
                table_lines.append(lines[index])
                index += 1
            rows = []
            for raw in table_lines:
                cells = [cell.strip() for cell in raw.strip("|").split("|")]
                if all(set(cell) <= {"-"} and cell for cell in cells):
                    continue
                rows.append(cells)
            if rows:
                parts.append("<table><thead><tr>")
                parts.extend(f"<th>{inline_html(cell)}</th>" for cell in rows[0])
                parts.append("</tr></thead><tbody>")
                for row in rows[1:]:
                    parts.append("<tr>" + "".join(f"<td>{inline_html(cell)}</td>" for cell in row) + "</tr>")
                parts.append("</tbody></table>")
            continue
        elif re.match(r"^\d+\.\s+", stripped):
            items = []
            while index < len(lines) and re.match(r"^\d+\.\s+", lines[index].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[index].strip()))
                index += 1
            parts.append("<ol>" + "".join(f"<li>{inline_html(item)}</li>" for item in items) + "</ol>")
            continue
        elif stripped.startswith("- "):
            items = []
            while index < len(lines) and lines[index].strip().startswith("- "):
                items.append(lines[index].strip()[2:])
                index += 1
            parts.append("<ul>" + "".join(f"<li>{inline_html(item)}</li>" for item in items) + "</ul>")
            continue
        else:
            css = " class=\"meta\"" if stripped.startswith(("Status:", "Date:", "Owner:", "Lane:", "Source handoff:")) else ""
            parts.append(f"<p{css}>{inline_html(stripped)}</p>")
        index += 1
    parts.extend(["</main></body>", "</html>"])
    return "\n".join(parts) + "\n"


def set_style_font(style, name: str, size: float | None = None, color: str | None = None) -> None:
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], fill_header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[cell_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0 and fill_header:
                set_cell_shading(cell, "E8EEF5")


def add_inline_runs(paragraph, value: str) -> None:
    chunks = re.split(r"(`[^`]+`)", value)
    for chunk in chunks:
        if not chunk:
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)
        else:
            paragraph.add_run(chunk)


def add_list(doc: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        add_inline_runs(paragraph, item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for col, header in enumerate(headers):
        paragraph = table.cell(0, col).paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        run.bold = True
    for row_index, row in enumerate(rows, 1):
        for col, value in enumerate(row):
            paragraph = table.cell(row_index, col).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline_runs(paragraph, value)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], fill_header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    first = True
    for line in text.splitlines():
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def build_docx(markdown: str, handoff: str, desk: dict) -> None:
    paste_block = extract_paste_block(handoff)
    counts = desk["receiver_handoff_readback"]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style_font(doc.styles["Normal"], "Calibri", 11, "172033")
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.25
    set_style_font(doc.styles["Heading 1"], "Calibri", 16, "2E74B5")
    set_style_font(doc.styles["Heading 2"], "Calibri", 13, "2E74B5")
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(18)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.text = "Harvey/Nerdkle Mack Paste Packet"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("5B6474")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Mack Architecture Tear-Apart Paste Packet")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    for line in [
        "Status: READY_TO_COPY_NOT_SENT",
        "Date: 2026-07-06",
        "Owner: Heimerdinker@Betsy",
        "Lane: Harvey/Nerdkle architecture review",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        meta = paragraph.add_run(line)
        meta.font.size = Pt(9.5)
        meta.font.color.rgb = RGBColor.from_string("5B6474")

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_geometry(callout, [9360], fill_header=False)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "F4F7FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "This is the copy surface. It does not send anything to Mack, does not claim Mack returned review, and does not create a next-build packet."
    )
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    doc.add_paragraph()

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "Generated from the live ready-to-paste handoff so Ben can review context, copy Mack's mission, and later capture Mack's return without inventing success."
    )

    doc.add_heading("Read Order For Ben", level=1)
    add_list(
        doc,
        [
            "Read `MACK_ARCHITECTURE_REVIEW_DESK_READOUT_20260706.md` first.",
            "Read `BEN_PRE_MACK_ARCHITECTURE_BRIEF_AND_AEYE_INPUT_20260706.md` next.",
            "Read `BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.md` third.",
            "Read `MACK_ARCHITECTURE_ATTACK_SCORECARD_20260706.md` fourth.",
            "Read the main architecture packet in Markdown, DOCX, or HTML.",
            "If Ben decides to send, copy the Mack block below to Mack.",
            "Paste Mack's returned `MACK REVIEW RETURN` block into the review intake and `MACK SCORECARD RETURN` into the scorecard intake, then run the validators.",
        ],
        numbered=True,
    )

    doc.add_heading("Current Proof Counts", level=1)
    add_table(
        doc,
        ["Count", "Value"],
        [
            ["Receiver handoffs indexed", str(counts["count"])],
            ["Posted", str(counts["posted_count"])],
            ["Pending receiver", str(counts["pending_count"])],
            ["Returned-unposted", str(counts["returned_unposted_count"])],
            ["Template-return blocked", str(counts["template_return_blocked_count"])],
            ["Invalid", str(counts["invalid_count"])],
            ["Malformed", str(counts["malformed_count"])],
        ],
        [3600, 5760],
    )

    doc.add_heading("Truth Boundary", level=1)
    add_list(
        doc,
        [
            "This packet does not send anything to Mack.",
            "Mack has not returned a review.",
            "No Mack receipt is claimed.",
            "No canonical next-build packet exists.",
            "Universal receiver proof is not claimed while pending receiver-return lanes remain.",
            "A transport ACK or blocked template is not receiver work completion.",
        ],
    )

    doc.add_heading("Copy/Paste Block For Mack", level=1)
    add_code_block(doc, paste_block)

    doc.add_heading("After Mack Returns", level=1)
    doc.add_paragraph(f"Paste Mack's returned `MACK REVIEW RETURN` block into `{repo_rel(RETURN_INTAKE)}`.")
    doc.add_paragraph("Then run `node scripts\\foreman\\mack-architecture-return-intake-validator.mjs`.")
    doc.add_paragraph(f"Paste Mack's returned `MACK SCORECARD RETURN` block into `{repo_rel(SCORECARD_INTAKE)}`.")
    doc.add_paragraph("Then run `node scripts\\foreman\\mack-architecture-scorecard-return-validator.mjs`.")
    doc.add_paragraph(
        "Only after Mack returns a complete structured block and Ben explicitly accepts the direction should `--convert --ben-accepted` be used."
    )

    doc.add_heading("Artifact Drawer", level=1)
    add_table(
        doc,
        ["Role", "Path"],
        [
            ["Desk readout MD", repo_rel(READOUT_MD)],
            ["Desk readout DOCX", repo_rel(READOUT_DOCX)],
            ["Desk readout HTML", repo_rel(READOUT_HTML)],
            ["Ben pre-Mack Aeye brief", repo_rel(BEN_BRIEF)],
            ["Connection map MD", repo_rel(CONNECTION_MD)],
            ["Connection map HTML", repo_rel(CONNECTION_HTML)],
            ["Connection map JSON", repo_rel(CONNECTION_JSON)],
            ["Connection map Mermaid", repo_rel(CONNECTION_MERMAID)],
            ["Mack attack scorecard MD", repo_rel(SCORECARD_MD)],
            ["Mack attack scorecard HTML", repo_rel(SCORECARD_HTML)],
            ["Mack attack scorecard JSON", repo_rel(SCORECARD_JSON)],
            ["Main packet MD", repo_rel(PACKET_MD)],
            ["Main packet DOCX", repo_rel(PACKET_DOCX)],
            ["Main packet HTML", repo_rel(PACKET_HTML)],
            ["Ready-to-paste handoff source", repo_rel(HANDOFF)],
            ["Return template", repo_rel(RETURN_TEMPLATE)],
            ["Return intake", repo_rel(RETURN_INTAKE)],
            ["Scorecard return intake", repo_rel(SCORECARD_INTAKE)],
        ],
        [2700, 6660],
    )

    footer = section.footer.paragraphs[0]
    footer.text = "Mack Architecture Tear-Apart Paste Packet - 2026-07-06"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("5B6474")

    doc.save(DOCX_PATH)


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main() -> None:
    ARCH_DIR.mkdir(parents=True, exist_ok=True)
    RECEIPT_DIR.mkdir(parents=True, exist_ok=True)

    handoff = HANDOFF.read_text(encoding="utf-8")
    desk = read_json(DESK_RECEIPT)
    markdown = build_markdown(handoff, desk)

    MD_PATH.write_text(markdown, encoding="utf-8")
    HTML_PATH.write_text(markdown_to_html(markdown), encoding="utf-8")
    build_docx(markdown, handoff, desk)

    docx_readback = docx_text(DOCX_PATH)
    generated_at = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    validation = {
        "markdown_written": MD_PATH.exists(),
        "html_written": HTML_PATH.exists(),
        "docx_written": DOCX_PATH.exists(),
        "source_handoff_ready_to_paste": "STATUS: READY_TO_PASTE" in handoff,
        "contains_readout_first_instruction": "Read the Mack review desk readout first" in handoff,
        "contains_connection_map_paths": "BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.md" in handoff
        and "BOOK_ARCHITECTURE_CONNECTION_MAP_20260706.md" in markdown,
        "contains_attack_scorecard_paths": "MACK_ARCHITECTURE_ATTACK_SCORECARD_20260706.md" in handoff
        and "MACK_ARCHITECTURE_ATTACK_SCORECARD_20260706.md" in markdown,
        "contains_return_contract": "MACK REVIEW RETURN" in handoff,
        "contains_scorecard_return_contract": "MACK SCORECARD RETURN" in handoff and "MACK SCORECARD RETURN" in markdown,
        "contains_scorecard_validator_command": "mack-architecture-scorecard-return-validator.mjs" in markdown,
        "mack_return_not_received_preserved": desk["validation"]["validator_reports_mack_not_received"] is True,
        "canonical_next_build_packet_absent": desk["validation"]["canonical_next_build_packet_absent"] is True,
        "external_send_not_claimed": desk["validation"]["no_external_send_claim"] is True,
        "universal_receiver_proof_not_claimed": "Do not claim universal receiver proof" in handoff,
        "docx_structural_markers_present": all(
            marker in docx_readback
            for marker in [
                "READY_TO_COPY_NOT_SENT",
                "Copy/Paste Block For Mack",
                "MACK REVIEW RETURN",
                "MACK SCORECARD RETURN",
                "Mack has not returned a review.",
            ]
        ),
    }

    receipt = {
        "schema": "MACK_ARCHITECTURE_TEAR_APART_PASTE_PACKET_RECEIPT",
        "status": "ARTIFACT",
        "timestamp": generated_at,
        "machine": "BETSY",
        "agent": "Heimerdinker@Betsy",
        "packet_id": "MACK_HARVEY_NERDKLE_ARCHITECTURE_TEAR_APART_20260706",
        "receipt_id": f"{BASE}_RECEIPT",
        "repo": str(ROOT),
        "command": "python scripts/foreman/build-mack-architecture-paste-packet.py",
        "validation": validation,
        "receiver_handoff_readback": desk["receiver_handoff_readback"],
        "source_receipts": [repo_rel(DESK_RECEIPT)],
        "file_hashes": [
            file_hash(MD_PATH),
            file_hash(HTML_PATH),
            file_hash(DOCX_PATH),
            file_hash(HANDOFF),
            file_hash(Path(__file__).resolve()),
        ],
        "truth_boundary": "This paste packet is a local generated copy surface. It does not send anything to Mack, claim Mack returned review, claim universal receiver proof, or generate a next-build packet.",
        "stop_conditions_respected": [
            "no external send",
            "no Mack receipt claim",
            "no universal receiver proof claim",
            "no next-build packet generated",
            "no deploy",
            "no push",
            "no secrets",
        ],
    }
    RECEIPT_PATH.write_text(json.dumps(receipt, indent=2) + "\n", encoding="utf-8")

    print(
        json.dumps(
            {
                "status": "ARTIFACT",
                "markdown": repo_rel(MD_PATH),
                "html": repo_rel(HTML_PATH),
                "docx": repo_rel(DOCX_PATH),
                "receipt": repo_rel(RECEIPT_PATH),
                "docx_structural_markers_present": validation["docx_structural_markers_present"],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
